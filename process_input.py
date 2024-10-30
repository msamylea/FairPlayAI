from PyPDF2 import PdfReader
from io import BytesIO
from docx import Document
import chardet

def process_input(input_data):
    """
    Process the input, which can be either a file path, raw text, or binary data.
    """
    if isinstance(input_data, str):
        if input_data.endswith('.pdf'):
            return route_pdf(input_data)
        elif input_data.endswith('.docx'):
            return route_docx(input_data)
        elif input_data.endswith('.doc'):
            return route_doc(input_data)
        elif input_data.endswith('.txt'):
            return route_txt(input_data)
        else:
            # Assume it's raw text if it doesn't match any file extension
            return input_data
    elif isinstance(input_data, bytes):
        # Handle uploaded file content
        try:
            # Attempt to detect the encoding
            detected = chardet.detect(input_data)
            encoding = detected['encoding'] or 'utf-8'
            return input_data.decode(encoding)
        except UnicodeDecodeError:
            # If decoding fails, treat it as binary data
            return handle_binary_data(input_data)
    else:
        return "Unsupported input type"

def handle_binary_data(data):
    """
    Handle binary data by attempting to parse it as different file types.
    """
    # Try to parse as PDF
    try:
        pdf = PdfReader(BytesIO(data))
        text = ""
        for page in pdf.pages:
            text += page.extract_text()
        return text
    except:
        pass

    # Try to parse as DOCX
    try:
        doc = Document(BytesIO(data))
        return "\n".join([para.text for para in doc.paragraphs])
    except:
        pass

    # If all attempts fail, return an error message
    return "Unable to process the uploaded file. Please ensure it's a valid PDF, DOCX, or text file."

def route_pdf(path):
    try:
        reader = PdfReader(path)
        policy = ""
        for page in reader.pages:
            policy += page.extract_text()
        
        # Join pages
        policy = " ".join(policy.split("\n"))
    except Exception as e:
        return {"error": f"Error loading PDF: {str(e)}"}

    return policy

def route_docx(path):
    try:
        doc = Document(path)
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        return {"error": f"Error loading DOCX: {str(e)}"}

def route_doc(path):
    try:
        doc = Document(path)
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        return {"error": f"Error loading DOC: {str(e)}"}

def route_txt(path):
    try:
        with open(path, 'rb') as f:
            raw_data = f.read()
            detected = chardet.detect(raw_data)
            encoding = detected['encoding'] or 'utf-8'
            return raw_data.decode(encoding)
    except Exception as e:
        return {"error": f"Error loading TXT: {str(e)}"}